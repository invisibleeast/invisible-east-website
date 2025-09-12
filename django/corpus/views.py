from django.views.generic import (DetailView, ListView, TemplateView, View)
from django.db.models.functions import Lower
from django.db.models import (Count, Q, CharField, TextField, Prefetch)
from django.http import HttpResponseRedirect, HttpResponse, JsonResponse
from django.urls import reverse
from django.conf import settings
from functools import reduce
from operator import (or_, and_)
from bs4 import BeautifulSoup
from docx import Document
from docx.shared import Cm
from . import models
import json
import math
import re
import os
import glob
import time


# Sections of this file:
# 1. Reusable Code
# 2. Main Views
# 3. Maps Views
# 4. Corpus Insights (data visualisations) Views
# 5. Download Data Views


#
# 1. Reusable Code
#


# Special starts to the values & labels of options in 'filter' select lists
filter_pre = 'filter_'
filter_pre_mm = f'{filter_pre}mm_'  # Many to Many relationship
filter_pre_fk = f'{filter_pre}fk_'  # Foreign Key relationship
filter_pre_gt = f'{filter_pre}gt_'  # Greater than (or equal to) filter, e.g. "Date (from)"
filter_pre_lt = f'{filter_pre}lt_'  # Less than (or equal to) filter, e.g. "Date (to)"
filter_pre_bl = f'{filter_pre}bl_'  # Boolean, e.g. "Has transcription"


def clean_date_from_datetime(datetime_obj):
    """
    Return a clean date in format DD/MM/YYYY (e.g. 31/01/2023) from a datetime object
    """
    return datetime_obj.strftime('%d/%m/%Y') if datetime_obj else None


def replace_nth(string, sub, wanted, nth):
    """
    Replace only the nth instance of a substring in a string
    """

    where = [m.start() for m in re.finditer(sub, string)][nth]
    before = string[:where]
    after = string[where:]
    after = after.replace(sub, wanted, 1)
    new_string = before + after

    return new_string


def html_details_link_to_text_list_filtered(filter_id, filter_object, link_text=None):
    """
    Returns a HTML anchor tag <a> that will link to the Corpus Text List page
    and filter on the object being clicked.
    E.g. if clicking 'Bactrian' in the detail then the user can see all other Bactrian Corpus Texts
    If link_text is not provided then it will automatically show the object's str as the link text
    """
    # If no link text is provided, use the object str by default
    if not link_text:
        link_text = filter_object
    # Return the complete HTML anchor tag
    return f'<a href="{reverse("corpus:text-list")}?{filter_id}={filter_object.id}">{link_text}</a>' if filter_object else None


def html_details_list_items(filter_id, object_list):
    """
    Return a HTML string of a list of objects (i.e. a queryset) for use in the 'Details' tab of an item page.
    E.g. showing ManyToMany and reverse FK objects
    """

    if len(object_list):
        object_links_list = [html_details_link_to_text_list_filtered(filter_id, obj) for obj in object_list]

        # Multiple objects
        if len(object_list) > 1:
            list_items = '</li><li>'.join(object_links_list)
            return f'<ul><li>{list_items}</li></ul>'
        # 1 object
        elif len(object_list) == 1:
            return object_links_list[0]

    # No objects
    return None


def request_post_get_safe(request, name):
    """
    Gets a value from a POST request
    Returns None if an empty value is found
    """

    value = request.POST.get(name)
    return value if value and len(value) else None


def text_initial_queryset(user):
    """
    Returns the initially filtered list of Text objects used in get_queryset() methods of views below, e.g. TextDetailView and TextListView
    """

    # If user is logged in, show all texts
    if user.is_authenticated:
        return models.Text.objects.all()
    # If user is not logged in, show only approved texts
    else:
        return models.Text.objects.filter(public_review_approved=True)


def media_url_full(request, media_file_path):
    """
    Return the full URL of the media file
    """
    return f"{request.scheme}://{request.META['HTTP_HOST']}{media_file_path}"


# Used in clean_html function, but compiled here once for performance improvements
CLEANR = re.compile('<.*?>')


def clean_html(raw_html):
    """
    Remove all tags from HTML and convert chars
    """
    clean_text = re.sub(CLEANR, '', str(raw_html))
    clean_text = clean_text\
        .replace('&#39;', "'")\
        .replace('&nbsp;', ' ')\
        .replace('&quot;', '"')
    return clean_text


def downloaddata_text_queryset(view_request):
    """
    Generates the queryset of Text objects for use in downloaddata views
    """
    texts = models.Text.objects.all().select_related(
        'primary_language__script',
        'type__category',
        'gregorian_date_century',
        'collection',
        'admin_principal_editor',
        'admin_classification',
        'corpus',
        'writing_support',
        'fold_lines_alignment',
        'admin_source_of_data',
        'document_subtype__category'
    ).prefetch_related(
        Prefetch(
            'texts',
            models.Text.objects.all().select_related('collection')
        ),
        'text_related_publications__publication',
        'text_dates',
        Prefetch(
            'additional_languages',
            models.SlTextLanguage.objects.all().select_related('script')
        ),
        Prefetch(
            'text_folios',
            models.TextFolio.objects.all().select_related('side', 'open_state')
        ),
        Prefetch(
            'persons_in_texts',
            models.PersonInText.objects.all().select_related('person_role_in_text', 'person__gender')
        ),
        'persons_in_texts__person__person_1__person_2',
        'persons_in_texts__person__person_1__relationship_type',
    )

    # Hide non-approved Text data from public users
    if not view_request.user.is_authenticated:
        texts = texts.filter(public_review_approved=True)

    # Limit to only specific texts, if provided
    text_ids = view_request.GET.get('text_ids', None)
    if text_ids:
        texts = texts.filter(id__in=text_ids.split(','))

    return texts


def json_str(value):
    """
    For use in JSON data download, convert value to string
    if value exists, else return None/null
    """
    return str(value) if value else None


#
# 2. Main Views
#


class TextDetailView(DetailView):
    """
    Class-based view for Text detail template
    """
    template_name = 'corpus/text-detail.html'
    model = models.Text

    def get_queryset(self):
        # Start with the initial queryset of Text objects
        queryset = text_initial_queryset(self.request.user)

        # Improve performance
        queryset = queryset.select_related(
            'primary_language__script',
            'type__category',
            'gregorian_date_century',
            'collection',
        ).prefetch_related(
            'text_folios__text_folio_tags',
            'texts',
            'text_related_publications',
            'text_dates',
            'seals',
            Prefetch(
                'toponyms__texts',
                models.Text.objects.all().select_related('collection', 'primary_language__script')
            ),
            Prefetch(
                'persons_in_texts',
                models.PersonInText.objects.all().select_related('person_role_in_text', 'person__gender')
            ),
            'persons_in_texts__person__persons',
            'persons_in_texts__person__person_1__person_2',
            'persons_in_texts__person__person_1__relationship_type',
        )

        return queryset

    def get_context_data(self, **kwargs):
        context = super().get_context_data(**kwargs)

        # Codex pagination
        if self.object.is_codex:
            codex_pagination = int(self.request.GET.get('codex_pagination', 0))
            codex_perpage = 25
            codex_images = self.object.codex_images
            context['codex_images'] = codex_images[
                codex_pagination:codex_pagination + codex_perpage
            ]
            context['codex_pagination_pagecountstart'] = codex_pagination
            # Codex pagination options (if there are multiple pages)
            if len(codex_images) > codex_perpage:
                codex_pages = range(0, math.ceil(len(codex_images) / codex_perpage))
                codex_pagination_options = []
                for p in codex_pages:
                    p_start = p * codex_perpage
                    p_end = min(p_start + codex_perpage, len(codex_images))
                    codex_pagination_options.append(
                        {'value': p_start, 'label': f'{p_start + 1} - {p_end}'}
                    )
                context['codex_pagination_options'] = codex_pagination_options

        # TextFolio and TextFolioTag
        context['text_folios'] = self.object.text_folios.all().select_related(
            'side',
            'open_state'
        ).prefetch_related(
            Prefetch(
                'text_folio_tags',
                models.TextFolioTag.objects.all().select_related('tag__category')
            )
        )
        context['text_folio_tag_categories'] = models.SlTextFolioTagCategory.objects.all().prefetch_related('tags__text_folio_tags')
        context['text_folio_tags'] = models.SlTextFolioTag.objects.all().select_related('category')
        context['permalink'] = self.request.build_absolute_uri().split('?')[0]

        # Details data
        context['data_items'] = [

            # Content
            {
                'section_header': 'Content',
                'section_header_fa': 'محتوا',
            },
            {
                'label': 'Summary of Content',
                'label_fa': 'خلاصه محتوا',
                'value': self.object.summary_of_content
            },

            # Dates (Gregorian and Original)
            {
                'section_header': 'Dates',
                'section_header_fa': 'تاریخ',
            },
            {
                'html': self.object.details_html_dates_full
            },

            # General
            {
                'section_header': 'Details',
                'section_header_fa': 'جزئیات',
            },
            {
                'label': 'Shelfmark/Title',
                'label_fa': 'شماره قفسه/عنوان',
                'value': self.object.shelfmark
            },
            {
                'label': 'Collection',
                'label_fa': 'مجموعه',
                'value': html_details_link_to_text_list_filtered(
                    f'{filter_pre_fk}collection',
                    self.object.collection
                )
            },
            {
                'label': 'Groups',
                'label_fa': '',
                'value': html_details_link_to_text_list_filtered(
                    f'{filter_pre_fk}corpus',
                    self.object.corpus
                )
            },
            {
                'label': 'Classification',
                'label_fa': 'طبقه‌بندی',
                'value': self.object.admin_classification.name_full if self.object.admin_classification else None
            },
            {
                'label': 'Primary Language',
                'label_fa': 'زبان اصلی',
                'value': html_details_link_to_text_list_filtered(
                    f'{filter_pre_fk}primary_language',
                    self.object.primary_language
                )
            },
            {
                'label': 'Additional Languages',
                'label_fa': '',
                'value': html_details_list_items(
                    f'{filter_pre_mm}additional_languages',
                    self.object.additional_languages.all()
                )
            },
            {
                'label': 'Type',
                'label_fa': 'نوع',
                'value': html_details_link_to_text_list_filtered(
                    f'{filter_pre_fk}type',
                    self.object.type,
                    self.object.type.name if self.object.type else None
                )
            },
            {
                'label': 'Document Subtype',
                'label_fa': 'زیربخش سند',
                'value': html_details_link_to_text_list_filtered(
                    f'{filter_pre_fk}document_subtype',
                    self.object.document_subtype
                )
            },
            {
                'label': 'Toponyms',
                'label_fa': 'نام‌های جغرافیایی',
                'value': html_details_list_items(
                    f'{filter_pre_mm}toponyms',
                    self.object.toponyms.all()
                )
            },

            # Physical Description
            {
                'section_header': 'Physical Description',
                'section_header_fa': 'توصیفات فیزیکی',
            },
            {
                'label': 'Writing Support',
                'label_fa': 'سطح نوشتار',
                'value': html_details_link_to_text_list_filtered(
                    f'{filter_pre_fk}writing_support',
                    self.object.writing_support
                )
            },
            {
                'label': 'Writing Support Details',
                'label_fa': 'توضیحات سطح نوشتاری',
                'value': html_details_list_items(
                    f'{filter_pre_mm}writing_support_details',
                    self.object.writing_support_details.all()
                )
            },
            {
                'label': 'Writing Support Notes',
                'label_fa': 'نکات سطح نوشتار',
                'value': self.object.writing_support_details_additional
            },
            {
                'label': 'Height (cm)',
                'label_fa': 'طول (سانتی متر)',
                'value': self.object.dimensions_height
            },
            {
                'label': 'Width (cm)',
                'label_fa': 'عرض (سانتی متر)',
                'value': self.object.dimensions_width
            },
            {
                'label': 'Fold Lines',
                'label_fa': 'رد تا',
                'value': self.object.fold_lines_count
            },
            {
                'label': 'Fold Lines',
                'label_fa': 'رد تا',
                'value': html_details_link_to_text_list_filtered(
                    f'{filter_pre_fk}fold_lines_alignment',
                    self.object.fold_lines_alignment
                )
            },
            {
                'label': 'Fold Lines',
                'label_fa': 'رد تا',
                'value': self.object.fold_lines_details
            },

            # People
            {
                'section_header': 'People',
                'section_header_fa': 'اشخاص',
            },
            {
                'html': self.object.details_html_person_in_text
            },

            # Publications
            {
                'section_header': 'Publications',
                'section_header_fa': 'نشریات مرتبط',
            },
            {
                'html': self.object.details_html_publications
            },

            # Related Shelfmarks
            {
                'section_header': 'Related Shelfmarks',
                'section_header_fa': 'شماره قفسههای مرتبط',
            },
            {
                'html': self.object.details_html_texts
            },

            # IEDC Data
            {
                'section_header': 'IEDC Data',
                'section_header_fa': 'داده‌های پایگاه دیجیتال شرق مکنون',
            },
            {
                'label': 'IEDC ID',
                'label_fa': 'شناسه‌ی پایگاه دیجیتال شرق مکنون',
                'value': self.object.id
            },
            {
                'label': 'Date Added to IE Corpus',
                'label_fa': 'تاریخ اضافه شدن به پایگاه دیجیتال شرق مکنون',
                'value': clean_date_from_datetime(self.object.meta_created_datetime)
            },
            {
                'label': 'Date Last Updated in IE Corpus',
                'label_fa': 'تاریخ آخرین بروزرسانی در پایگاه شرق مکنون',
                'value': clean_date_from_datetime(self.object.meta_lastupdated_datetime)
            },

            # Citations
            {
                'section_header': 'Citations',
                'section_header_fa': 'ارجاعات',
            },
            {
                'label': 'Principal Editor',
                'label_fa': 'ویراستار اصلی',
                'value': self.object.admin_principal_editor
            },
            {
                'label': 'Contributors',
                'label_fa': 'مشارکت کنندگان',
                'value': self.object.admin_contributors_list
            },
            {
                'label': 'Source of Data',
                'label_fa': 'منبع دادهها',
                'value': self.object.admin_source_of_data
            },
            {
                'label': 'Suggested Citation',
                'label_fa': 'ارجاع پیشنهادی',
                'value': f'<a href="{reverse("general:about-cite")}">See \'How to Cite\'</a>'
            },
            {
                'label': 'Permalink',
                'label_fa': 'پیوند سند',
                'value': f'<a href="{context["permalink"]}">{context["permalink"]}</a>'
            },
            {
                'label': 'Image Permission Statement',
                'label_fa': 'بیانیه‌ی اجازه استفاده از تصویر',
                'value': self.object.image_permission_statement
            },

            # Contact
            {
                'section_header': 'Contact',
                'section_header_fa': 'تماس',
            },
            {
                'label': 'Contact Editorial Team',
                'label_fa': 'تماس با گروه تدوین',
                'value': f'<a href="mailto:{settings.MAIN_CONTACT_EMAIL}?subject=Invisible East Digital Corpus&body=This email relates to Text {self.object.id} - {context["permalink"]}">{settings.MAIN_CONTACT_EMAIL}</a> <em>(Please include the above permalink when contacting the editorial team about this Text)</em>'
            },

        ]

        return context


class TextListView(ListView):
    """
    Class-based view for Text list template
    """
    template_name = 'corpus/text-list.html'
    model = models.Text
    paginate_by = 30

    # Special starts to the values & labels of options in 'sort by' select lists, used in below sort() function and within views scripts
    sort_pre_count_value = 'count_'
    sort_pre_count_label = 'Number of '

    def get_field_type(self, field_name, queryset):
        """
        Return the type of a field
        E.g. used in sort() to see if case insensitivity is needed (if field is a CharField/TextField)
        """
        try:
            stripped_field_name = field_name.lstrip('-')
            if stripped_field_name in queryset.query.annotations:
                return queryset.query.annotations[stripped_field_name].output_field
            return queryset.model._meta.get_field(stripped_field_name)
        except Exception:
            return CharField  # If it fails, assume it's a CharField by default

    def get_queryset(self):
        # Start with the initial queryset of Text objects
        queryset = text_initial_queryset(self.request.user)

        # Improve performance
        queryset = queryset.select_related(
            'writing_support',
            'primary_language__script',
            'type__category',
            'gregorian_date_century',
            'collection',
        ).prefetch_related(
            'text_folios',
            'text_folios__text_folio_tags'
        )

        # Search
        try:
            searches = json.loads(self.request.GET.get('search', '[]'))
        except json.decoder.JSONDecodeError:
            searches = []

        field_names_to_search = [
            'shelfmark',
            'collection__name',
            'corpus__name',
            'primary_language__name',
            'primary_language__script__name',
            'type__name',
            'summary_of_content',
            'text_folios__transcription',
            'text_folios__translation',
            'text_folios__transliteration',
            'text_folios__text_folio_tags__tag__name'
        ]
        # Set list of search options
        if searches not in [[''], []]:
            search_type = 'iregex' if self.request.GET.get('search_type', '') == 'regex' else 'icontains'
            operator = or_ if self.request.GET.get('search_operator', '') == 'or' else and_
            queries = []
            for search in searches:
                # Uses 'or_' as the search term could appear in any field, so 'and_' wouldn't be suitable
                queries.append(reduce(or_, (Q((f'{field_name}__{search_type}', search.strip())) for field_name in field_names_to_search)))
            # Connect the individual search queries via the user-defined operator (or_ / and_)
            queries = reduce(operator, queries)
            # Filter the queryset using the completed search query
            queryset = queryset.filter(queries)

        # Filter
        for filter_key in [k for k in list(self.request.GET.keys()) if k.startswith(filter_pre)]:
            filter_value = self.request.GET.get(filter_key, '')
            # Perform the filter based on the type
            if filter_value != '':
                # Many to Many relationship (uses __in comparison and filter_value is a list)
                if filter_key.startswith(filter_pre_mm):
                    filter_field = filter_key.replace(filter_pre_mm, '')
                    queryset = queryset.filter(**{f'{filter_field}__in': [filter_value]})
                # Foreign Key relationship
                elif filter_key.startswith(filter_pre_fk):
                    filter_field = filter_key.replace(filter_pre_fk, '')
                    queryset = queryset.filter(**{filter_field: filter_value})
                # Greater than or equal to
                elif filter_key.startswith(filter_pre_gt):
                    filter_field = filter_key.replace(filter_pre_gt, '')
                    queryset = queryset.filter(**{f'{filter_field}__gte': filter_value})
                # Less than or equal to
                elif filter_key.startswith(filter_pre_lt):
                    filter_field = filter_key.replace(filter_pre_lt, '')
                    queryset = queryset.filter(**{f'{filter_field}__lte': filter_value})
                # Boolean content (e.g. field is not null or empty string)
                elif filter_key.startswith(filter_pre_bl):
                    filter_field = filter_key.replace(filter_pre_bl, '')
                    if filter_value == 'on':
                        try:
                            queryset = queryset.exclude(**{f'{filter_field}__isnull': True})  # remove null values
                        except Exception:
                            pass
                        try:
                            queryset = queryset.exclude(**{f'{filter_field}__exact': ''})  # remove empty strings
                        except Exception:
                            pass

        # Sort
        # Establish the sort direction (asc/desc) and the field to sort by, from the self.request
        sort = self.request.GET.get('sort', 'id')
        sort_dir = '-' if sort.startswith('-') else ''
        sort_pre_length = len(f"{sort_dir}{self.sort_pre_count_value}")  # e.g. '-numerical_' for descending numerical
        # Count sorting (e.g. sort by count of related items)
        if sort.startswith(self.sort_pre_count_value) or sort.startswith(f'-{self.sort_pre_count_value}'):
            # '-countitems' if descending, 'countitems' if ascending
            order_by = f'{sort_dir}countitems'
            sort_field = sort[sort_pre_length:]
            queryset = queryset.annotate(countitems=Count(sort_field)).order_by(order_by)
        # Standard sort
        else:
            # Sort descending (Z-A)
            if sort.startswith('-'):
                # Convert CharField and TextField values to lowercase, for case insensitivity
                if isinstance(self.get_field_type(sort, queryset), (CharField, TextField)):
                    queryset = queryset.order_by(Lower(sort[1:]).desc())
                else:
                    queryset = queryset.order_by(sort)
            # Sort ascending (A-Z)
            else:
                # Convert CharField and TextField values to lowercase, for case insensitivity
                if isinstance(self.get_field_type(sort, queryset), (CharField, TextField)):
                    queryset = queryset.order_by(Lower(sort))
                else:
                    queryset = queryset.order_by(sort)

        # return queryset, removing duplicates
        return queryset.distinct()

    def get_context_data(self, **kwargs):
        context = super().get_context_data(**kwargs)

        # Count of all published texts
        context['count_all_texts'] = text_initial_queryset(self.request.user).count()
        # Filter pre values
        context['filter_pre'] = filter_pre
        context['filter_pre_gt'] = filter_pre_gt
        context['filter_pre_lt'] = filter_pre_lt
        context['filter_pre_bl'] = filter_pre_bl

        # Download Data
        if self.request.GET:
            context['downloaddata_text_ids'] = ','.join([str(i) for i in self.object_list.values_list('id', flat=True)])

        # Results count start
        if context.get('is_paginated'):
            page_obj = context['page_obj']
            start_count = (page_obj.number - 1) * page_obj.paginator.per_page + 1
            context['start_count'] = start_count
        else:
            context['start_count'] = 1

        # Tag
        tag_id = self.request.GET.get(f'{filter_pre_fk}text_folios__text_folio_tags__tag', None)
        if tag_id:
            context['filter_active_tag'] = models.SlTextFolioTag.objects.get(id=tag_id).name

        # Options: Sort
        context['options_sort'] = [
            {
                'value': 'id',
                'label': 'Random',
                'label_fa': 'اتفاقی'
            },
            {
                'value': 'shelfmark',
                'label': 'Shelfmark (A-Z)',
                'label_fa': 'شماره قفسه (الفبایی)'
            },
            {
                'value': 'gregorian_date_century__century_number',
                'label': 'Converted date (CE) ↑',
                'label_fa': 'تاریخ میلادی ↑'
            },
            {
                'value': '-gregorian_date_century__century_number',
                'label': 'Converted date (CE) ↓',
                'label_fa': 'تاریخ میلادی ↓'
            },
            {
                'value': 'meta_created_datetime',
                'label': 'Date added to IEDC ↑',
                'label_fa': 'تاریخ اضافه شدن به پایگاه ↑'
            },
            {
                'value': '-meta_created_datetime',
                'label': 'Date added to IEDC ↓',
                'label_fa': 'تاریخ اضافه شدن به پایگاه ↓'
            }
        ]

        # Reused querysets in below filters (specified here to avoid duplicate SQL queries)
        filter_queryset_centuries = models.SlTextGregorianCentury.objects.all()

        # Includes (aka checkbox filters)
        context['options_includes'] = [
            {
                'filter_id': f'{filter_pre_bl}text_folios__transliteration',
                'filter_name': 'Transliteration',
                'filter_name_fa': 'رونویسی',
                'filter_helptext': 'a rendition of a text in another script, preserving pronunciation'
            },
            {
                'filter_id': f'{filter_pre_bl}text_folios__transcription',
                'filter_name': 'Transcription',
                'filter_name_fa': 'نسخه‌برداری ',
                'filter_helptext': 'reproduce spoken word in writing'
            },
            {
                'filter_id': f'{filter_pre_bl}text_folios__translation',
                'filter_name': 'Translation',
                'filter_name_fa': 'ترجمه'
            },
            {
                'filter_id': f'{filter_pre_bl}text_folios__image',
                'filter_name': 'Image',
                'filter_name_fa': 'عکس '
            },
            {
                'filter_id': f'{filter_pre_bl}text_folios__palaeography',
                'filter_name': 'Palaeography',
                'filter_name_fa': 'پالئوگرافی'
            },
            {
                'filter_id': f'{filter_pre_bl}seals__id',
                'filter_name': 'Seal',
                'filter_name_fa': 'مهر'
            },
            {
                'filter_id': f'{filter_pre_bl}codex_images_location',
                'filter_name': 'Codex',
                'filter_name_fa': 'کتاب خطی',
            }
        ]

        # Gregorian Dates
        context['options_datefilters'] = [
            {
                'filter_id': f'{filter_pre_gt}gregorian_date_century__century_number',
                'filter_classes': filter_pre_gt,
                'filter_name': 'From',
                'filter_name_fa': 'از',
                'filter_options': filter_queryset_centuries
            },
            {
                'filter_id': f'{filter_pre_lt}gregorian_date_century__century_number',
                'filter_classes': filter_pre_lt,
                'filter_name': 'To',
                'filter_name_fa': 'به',
                'filter_options': filter_queryset_centuries
            },
        ]

        # Filters
        context['options_filters'] = [
            {
                'filter_id': f'{filter_pre_fk}text_folios__text_folio_tags__tag',
                'filter_name': 'Tag',
                'filter_options': models.SlTextFolioTag.objects.all(),
                'filter_hidden': True  # this filter is not visible as its set by user clicking tags in list items
            },
            {
                'filter_id': f'{filter_pre_fk}primary_language',
                'filter_name': 'Primary Language',
                'filter_name_fa': 'زبان اصلی',
                'filter_options': models.SlTextLanguage.objects.filter(texts_primary__isnull=False).distinct().select_related('script')
            },
            {
                'filter_id': f'{filter_pre_mm}additional_languages',
                'filter_name': 'Additional Languages',
                'filter_name_fa': 'زبانهای دیگر',
                'filter_options': models.SlTextLanguage.objects.filter(texts__isnull=False).distinct().select_related('script')
            },
            {
                'filter_id': f'{filter_pre_fk}collection',
                'filter_name': 'Collection',
                'filter_name_fa': 'مجموعه',
                'filter_options': models.SlTextCollection.objects.all()
            },
            {
                'filter_id': f'{filter_pre_fk}corpus',
                'filter_name': 'Groups',
                'filter_name_fa': 'گروه',
                'filter_options': models.SlTextCorpus.objects.all(),
                'info_alert': 'These are sub-corpora of documents organised by place of origin (presumed or confirmed)',
                'info_alert_fa': 'این زیرمجموعه‌ها شامل اسنادی هستند که بر اساس محل نگارش (تأیید شده یا مفروض) طبقه‌بندی شده‌اند'
            },
            {
                'filter_id': f'{filter_pre_fk}type',
                'filter_name': 'Type of Text',
                'filter_name_fa': 'نوع متن',
                'filter_options': models.SlTextType.objects.all().select_related('category'),
                'info_link': reverse('general:about-typology')
            },
            {
                'filter_id': f'{filter_pre_fk}writing_support',
                'filter_name': 'Writing Support',
                'filter_name_fa': 'سطح نوشتار',
                'filter_options': models.SlTextWritingSupport.objects.all()
            },
        ]

        return context


class TextFolioTagCreateView(View):
    """
    Class-based view to create a TextFolioTag object in the database
    """

    def post(self, request):
        """
        Process the TextFolioTag object and related data objects
        """

        fail_response = HttpResponseRedirect(reverse('corpus:textfoliotag-failed'))

        try:

            text_folio_id = request_post_get_safe(request, 'textfolio')
            text_folio_obj = models.TextFolio.objects.get(id=text_folio_id)

            # Get an existing TextFolioTag object to link to a piece of trans text
            link_trans_text_to_tag = request_post_get_safe(request, 'linktranstexttotag')
            link_trans_text_to_tag = json.loads(link_trans_text_to_tag) if link_trans_text_to_tag else None
            if link_trans_text_to_tag and 'textFolioTagExistingId' in link_trans_text_to_tag:
                text_folio_tag = models.TextFolioTag.objects.get(id=int(link_trans_text_to_tag['textFolioTagExistingId']))

            else:
                # Get/Create SlTextFolioTag object
                tag_category = request_post_get_safe(request, 'category')
                tag_existing_id = request_post_get_safe(request, 'tag_existing')
                tag_new_name = request_post_get_safe(request, 'tag_new')
                if tag_existing_id:
                    tag = models.SlTextFolioTag.objects.get(id=tag_existing_id)
                elif tag_new_name:
                    tag = models.SlTextFolioTag.objects.create(
                        name=tag_new_name,
                        category=models.SlTextFolioTagCategory.objects.get(id=tag_category)
                    )
                else:
                    return fail_response

                # Create a TextFolioTag object (if a TextFolio obj is available)
                if text_folio_obj:
                    text_folio_tag = models.TextFolioTag.objects.create(
                        text_folio=text_folio_obj,
                        tag=tag,
                        details=request_post_get_safe(request, 'details'),
                        image_part_left=request_post_get_safe(request, 'image_part_left'),
                        image_part_top=request_post_get_safe(request, 'image_part_top'),
                        image_part_width=request_post_get_safe(request, 'image_part_width'),
                        image_part_height=request_post_get_safe(request, 'image_part_height'),
                        meta_created_by=request.user
                    )
                else:
                    return fail_response

            # Set the link to this text folio tag in the trans text, if this data exists
            if link_trans_text_to_tag:
                trans_text = getattr(text_folio_obj, link_trans_text_to_tag['textTrans'])
                trans_text_lines = trans_text.split('</li>')
                trans_text_line = trans_text_lines[int(link_trans_text_to_tag['textTransLineIndex'])]
                # Replaces only the required instance of the selected text (in case multiple instances)
                trans_text_line_new = replace_nth(
                    trans_text_line,
                    link_trans_text_to_tag['textSelected'],
                    f'<var data-textfoliotag="{text_folio_tag.id}">{link_trans_text_to_tag["textSelected"]}</var>',
                    int(link_trans_text_to_tag['textSelectedInstanceCountInLine'])
                )
                trans_text_lines[int(link_trans_text_to_tag['textTransLineIndex'])] = trans_text_line_new
                trans_text_new = '</li>'.join(trans_text_lines)
                # Update the TextFolio object with this new link in the trans text
                setattr(text_folio_obj, link_trans_text_to_tag['textTrans'], trans_text_new)
                text_folio_obj.save()

            # Return the user to the current page
            return HttpResponseRedirect(f"{reverse('corpus:text-detail', args=[request.POST.get('text')])}?tab=tags&textfolio={text_folio_id}")

        except Exception as e:
            print(e)
            return fail_response


class TextFolioTagFailedTemplateView(TemplateView):
    """
    Class based view to show a template that tells user the attempt to save the TextFolioTag failed
    """

    template_name = 'corpus/textfoliotag-failed.html'


class TextFolioTransLineDrawnOnImageManageView(View):
    """
    Class-based view to manage (create/edit/delete) a drawing of a line of trans text from a text folio on an image
    """

    def post(self, request):
        """
        Process the drawing on image data for the specified trans field in the TextFolio object
        """

        fail_response = HttpResponseRedirect(reverse('corpus:textfoliotranslinedrawnonimage-failed'))

        try:
            # Get data from request
            text = request_post_get_safe(request, 'text')
            trans_field = request_post_get_safe(request, 'trans_field')
            line_index = int(request_post_get_safe(request, 'line_index'))
            image_part_left = request_post_get_safe(request, 'image_part_left')
            image_part_top = request_post_get_safe(request, 'image_part_top')
            image_part_width = request_post_get_safe(request, 'image_part_width')
            image_part_height = request_post_get_safe(request, 'image_part_height')
            delete_image_part = request_post_get_safe(request, 'delete_image_part')

            # Get the TextFolio object and specified trans field
            text_folio_id = request_post_get_safe(request, 'textfolio')
            text_folio_obj = models.TextFolio.objects.get(id=text_folio_id)
            text_folio_obj_trans_field = getattr(text_folio_obj, trans_field)

            # Use BS to add data attributes to the chosen li element (i.e. the line of text being drawn)
            text_folio_obj_trans_field_html = BeautifulSoup(text_folio_obj_trans_field, features="html.parser")

            # Manage (add/edit/delete) the data for this trans line
            trans_line = text_folio_obj_trans_field_html.find_all(['li', 'td'])[line_index]
            # If the image part drawing is being deleted submitDrawLineOnImageForm(deleteImagePartDrawing=false);
            if delete_image_part:
                del trans_line['data-imagepartleft']
                del trans_line['data-imageparttop']
                del trans_line['data-imagepartwidth']
                del trans_line['data-imagepartheight']
            # If the image part drawing is being added/edited
            else:
                trans_line['data-imagepartleft'] = image_part_left
                trans_line['data-imageparttop'] = image_part_top
                trans_line['data-imagepartwidth'] = image_part_width
                trans_line['data-imagepartheight'] = image_part_height

            setattr(text_folio_obj, trans_field, str(text_folio_obj_trans_field_html))
            text_folio_obj.save()

            # Return the user to the current page
            return HttpResponseRedirect(f"{reverse('corpus:text-detail', args=[text])}?tab={trans_field}&textfolio={text_folio_id}")

        except Exception as e:
            print(e)
            return fail_response


class TextFolioTransLineDrawnOnImageFailedTemplateView(TemplateView):
    """
    Class based view to show a template that tells user the attempt to save the drawing of a line of text in a TextFolio failed
    """

    template_name = 'corpus/textfoliotranslinedrawnonimage-failed.html'


#
# 3. Maps Views
#


class MapTextsListView(ListView):
    """
    Class based view to show a map (of SlTextToponym objects) list template
    """

    template_name = 'corpus/map-iedctoponyms.html'
    model = models.SlTextToponym

    def get_queryset(self):
        # Start with the initial queryset of SlTextToponym objects that have coordinates data
        queryset = self.model.objects.filter(latitude__isnull=False, longitude__isnull=False)

        # Improve performance
        queryset = queryset.prefetch_related(
            Prefetch(
                'texts',
                models.Text.objects.all().select_related(
                    'primary_language',
                    'collection',
                )
            )
        )

        return queryset


class MapFindSpotTemplateView(TemplateView):
    """
    Class based view to show a map (of find spots) template
    """

    template_name = 'corpus/map-findspots.html'


#
# 4. Corpus Insights (data visualisations) Views
#


class InsightsLanguagesTemplateView(TemplateView):
    """
    Class based view to show a 'corpus insights - languages' template
    """

    template_name = 'corpus/insights-languages.html'

    def get_context_data(self, **kwargs):
        context = super().get_context_data(**kwargs)

        context['data'] = 1

        return context


#
# 5. Download Data Views
#


def downloaddata_word(request):
    """
    Creates a Word Document (.docx) containing Text (and related) data
    and returns the file to be download
    """

    texts = downloaddata_text_queryset(request)
    texts_count = texts.count()

    # Create a new Document
    word_doc = Document()

    # Set page margins
    sections = word_doc.sections
    margin = Cm(1.5)
    for section in sections:
        section.top_margin = margin
        section.bottom_margin = margin
        section.left_margin = margin
        section.right_margin = margin

    # Build content within Document:

    # Introduction content
    word_doc.add_heading('Invisible East Digital Corpus', 0)
    word_doc.add_paragraph(f"This document contains data for {texts_count} text{'s' if texts_count != 1 else ''} exported from the IEDC database.")
    word_doc.add_paragraph(f"The data was exported on {time.strftime('%d/%m/%Y')} at {time.strftime('%H:%M')}.")
    word_doc.add_paragraph(f"If you have any questions or comments please contact the IEDC Editorial Team: {settings.MAIN_CONTACT_EMAIL}")

    # Content for each text
    for i, text in enumerate(texts):

        permalink = f"{request.build_absolute_uri('/')[:-1]}{reverse('corpus:text-detail', args=[text.id])}"

        word_doc.add_page_break()
        heading = f'Text {i + 1} of {texts_count}: {text.title}' if texts_count > 1 else str(text.title)
        word_doc.add_heading(heading, 1)
        word_doc.add_paragraph(f'Source: {permalink}')

        text_data_items = [

            # General
            {'section_header': 'Details'},
            {'label': 'Shelfmark/Title', 'value': text.shelfmark},
            {'label': 'Collection', 'value': text.collection},
            {'label': 'Group', 'value': text.corpus},
            {'label': 'Classification', 'value': text.admin_classification.name_full if text.admin_classification else None},
            {'label': 'Primary Language', 'value': text.primary_language},
            {'label': 'Additional Languages', 'value': text.strlist_additional_languages},
            {'label': 'Type', 'value': text.type.name if text.type else None},
            {'label': 'Document Subtype', 'value': text.document_subtype},
            {'label': 'Toponyms', 'value': text.strlist_toponyms},

            # Physical Description
            {'section_header': 'Physical Description'},
            {'label': 'Writing Support', 'value': text.writing_support},
            {'label': 'Writing Support Details', 'value': text.strlist_writing_support_details},
            {'label': 'Writing Support Notes', 'value': text.writing_support_details_additional},
            {'label': 'Height (cm)', 'value': text.dimensions_height},
            {'label': 'Width (cm)', 'value': text.dimensions_width},
            {'label': 'Fold Lines Count', 'value': text.fold_lines_count},
            {'label': 'Fold Lines Alignment', 'value': text.fold_lines_alignment},
            {'label': 'Fold Lines', 'value': text.fold_lines_details},

            # Content
            {'section_header': 'Content Summary'},
            {'content_block': text.summary_of_content},

            # Dates (Gregorian and Original)
            {'section_header': 'Dates'},
            {'content_block': text.gregorian_date_full},

            # People
            {'section_header': 'People'},
            {'content_block': text.strlist_persons_in_texts},

            # Publications
            {'section_header': 'Publications'},
            {'content_block': text.strlist_publications},

            # Related Shelfmarks
            {'section_header': 'Related Shelfmarks'},
            {'content_block': text.strlist_texts},

            # IEDC Data
            {'section_header': 'IEDC Data'},
            {'label': 'IEDC ID', 'value': text.id},
            {'label': 'Date Added to IE Corpus', 'value': clean_date_from_datetime(text.meta_created_datetime)},
            {'label': 'Date Last Updated in IE Corpus', 'value': clean_date_from_datetime(text.meta_lastupdated_datetime)},

            # Citations
            {'section_header': 'Citations'},
            {'label': 'Principal Editor', 'value': text.admin_principal_editor},
            {'label': 'Contributors', 'value': text.admin_contributors_list},
            {'label': 'Source of Data', 'value': text.admin_source_of_data},
            {'label': 'Permalink', 'value': permalink},
            {'label': 'Image Permission Statement', 'value': clean_html(text.image_permission_statement)},

            # Folios
            {'section_header': 'Folios'},
            {'content_block': text.strlist_text_folios},

            # Transliteration
            {'section_header': 'Transliteration'},
            {'content_block': text.transliteration_text_lines_str},

            # Transcription
            {'section_header': 'Transcription'},
            {
                'content_block': text.transcription_text_lines_str,
                'alignment': 2 if text.primary_language.script.is_written_right_to_left else 0
            },

            # Transcription
            {'section_header': 'Translation'},
            {'content_block': text.translation_text_lines_str},
        ]

        for text_data_item in text_data_items:
            # Section headers
            if 'section_header' in text_data_item:
                p = word_doc.add_paragraph()
                section_header = p.add_run(f"\n{text_data_item['section_header']}\n")
                section_header.underline = True
            # Label/Value keypairs
            elif 'label' in text_data_item and 'value' in text_data_item and text_data_item['value']:
                p = word_doc.add_paragraph()
                # Label
                label = p.add_run(f"{text_data_item['label']}: ")
                label.bold = True
                # Value
                p.add_run(str(text_data_item['value']))
            # Content blocks (e.g. multiline blocks of text, like transcriptions)
            elif 'content_block' in text_data_item and text_data_item['content_block']:
                p = word_doc.add_paragraph()
                p.add_run(f"{clean_html(text_data_item['content_block'])}")
                if 'alignment' in text_data_item:
                    p.alignment = text_data_item['alignment']

            # Add formatting to p
            paragraph_format = p.paragraph_format
            paragraph_format.space_before = Cm(0)
            paragraph_format.space_after = Cm(0)

    # Delete all existing files in the export directory
    data_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'downloaddata', 'word')
    files = glob.glob(data_path + '/*')
    for f in files:
        os.remove(f)

    # Establish new file name
    file_name = f'iedc_{time.strftime("%Y-%m-%d_%H-%M")}.docx'
    file_path = os.path.join(data_path, file_name)

    # Save Word document and return the file
    word_doc.save(file_path)
    if os.path.exists(file_path):
        with open(file_path, 'rb') as fh:
            response = HttpResponse(fh.read(), content_type="application/word")
            response['Content-Disposition'] = f'inline; filename={os.path.basename(file_path)}'
            return response


def downloaddata_json(request):
    """
    Returns a JSON object containing all Text and related data
    """

    texts = downloaddata_text_queryset(request)
    data = []
    for text in texts:
        permalink = f"{request.build_absolute_uri('/')[:-1]}{reverse('corpus:text-detail', args=[text.id])}"

        data.append({
            # General
            'uri': json_str(text.uri),
            'shelfmark': json_str(text.shelfmark),
            'collection': json_str(text.collection),
            'group': json_str(text.corpus),
            'classification': text.admin_classification.name_full if text.admin_classification else None,
            'primaryLanguage': json_str(text.primary_language),
            'additionalLanguages': [str(lng) for lng in text.additional_languages.all()],
            'documentType': text.type.name if text.type else None,
            'documentSubtype': json_str(text.document_subtype),
            'toponyms': [
                {
                    'name': t.name,
                    'alternativeReadings': t.alternative_readings,
                    'otherAttestedForms': t.other_attested_forms,
                    'latitude': t.latitude,
                    'longitude': t.longitude,
                    'urls': t.urls
                } for t in text.toponyms.all()
            ],

            # Physical Description
            'writingSupport': json_str(text.writing_support),
            'writingSupportDetails': [str(w) for w in text.writing_support_details.all()],
            'writingSupportNotes': text.writing_support_details_additional,
            'height': json_str(text.dimensions_height),
            'width': json_str(text.dimensions_width),
            'foldLinesCount': json_str(text.fold_lines_count),
            'foldLinesAlignment': json_str(text.fold_lines_alignment),
            'foldLines': json_str(text.fold_lines_details),

            # Content
            'contentSummary': text.summary_of_content,

            # Dates (Gregorian and Original)
            'dates': text.gregorian_date_full,

            # Folios
            'folios': [
                {
                    'side': json_str(f.side),
                    'openState': json_str(f.open_state),
                    'image': f.image.url if f.image else None,
                    'transliteration': f.transliteration,
                    'transcription': f.transcription,
                    'translation': f.translation,
                    'palaeography': f.palaeography
                } for f in text.text_folios.all()
            ],

            # People
            'personsInText': [
                {
                    'person': json_str(p.person),
                    'personNameInText': p.person_name_in_text,
                    'personRoleInText': json_str(p.person_role_in_text)
                } for p in text.persons_in_texts.all()
            ],

            # Publications
            'publications': [
                {
                    'publication': json_str(p.publication),
                    'pages': p.pages,
                    'catalogueNumber': p.catalogue_number,
                    'details': p.details
                } for p in text.text_related_publications.all()
            ],

            # Related Shelfmarks
            'relatedShelfmarks': [
                {
                    'id': t.id,
                    'title': json_str(t),
                } for t in text.texts.all()
            ],

            # IEDC Data
            'iedcId': text.id,
            'dateAdded': clean_date_from_datetime(text.meta_created_datetime),
            'dateLastUpdated': clean_date_from_datetime(text.meta_lastupdated_datetime),

            # Citations
            'principalEditor': json_str(text.admin_principal_editor),
            'contributors': json_str(text.admin_contributors_list),
            'sourceOfData': json_str(text.admin_source_of_data),
            'permalink': permalink,
            'imagePermissionStatement': clean_html(text.image_permission_statement),
        })

    return JsonResponse(data, safe=False)
